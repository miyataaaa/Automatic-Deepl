from docx import Document
from docx.oxml.ns import qn
import os 
from webdriver_manager.chrome import ChromeDriverManager
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.keys import Keys
import requests
import time
from threading import Lock
import time
from concurrent.futures import ThreadPoolExecutor
import math 
import re
from copy import deepcopy
options = Options()
options.add_argument('--disable-gpu')
options.add_argument('--headless')
options.add_argument('--no-sandbox')
options.add_argument('--disable-extensions')
options.add_argument('--proxy-server="direct://"')
options.add_argument('--proxy-bypass-list=*')
options.add_argument('--start-maximized')
lock = Lock()

class auto_translator:
    
    def __init__(self, wating_sec: int=2, **kwargs):
        self.dirpath = kwargs['Fpath']
        self.fname = kwargs['Fname']
        self.fpath = os.path.join(self.dirpath, self.fname)
        self.docxfile = Document(self.fpath)
        self.tables = self.docxfile.tables[0]
        self.insert_pt = self._search_insertPt()
        self.eg_sentence = self._sentence_from_docx()
        self.translated_sentence = []
        self.max_worker = kwargs['max_worker']
        self.replace_dict_jp_to_jp = kwargs['replace_dict_jp_to_jp'] if len(kwargs['replace_dict_jp_to_jp']) != 0 else None
        self.replace_flag_jp_to_jp = True if type(self.replace_dict_jp_to_jp) == dict else False
        self.replace_dict_eg_to_jp = kwargs['replace_dict_eg_to_jp'] if len(kwargs['replace_dict_eg_to_jp']) != 0 else None
        self.replace_flag_eg_to_jp = True if type(self.replace_dict_eg_to_jp) == dict else False
        self.japaneseFont = kwargs['japaneseFont']
        self.wating_sec = wating_sec
        
        
    def _search_insertPt(self):
        
        """
        翻訳開始位置（insert_pt)を探索するメソッド。
        全ての行を探索して英文、和文の両方が入っている最後の行の位置を返す。
        """
        column_eg = self.tables.columns[0]
        column_jp = self.tables.columns[1]                    
        insert_pt = 0
        for i in range(len(column_eg.cells)):
            eg = column_eg.cells[i].text
            jp = column_jp.cells[i].text
            if (eg != "") and (jp != ""):
                insert_pt = i
        
        return insert_pt
    
    def _sentence_from_docx(self):

        """
        翻訳対象のPDFファイルから、翻訳したい部分を手動でコピペしたwordファイルから英文を抜き出して
        1つのリストにする関数。
        ファイル内の構造は1つの2列×n行テーブルを持ち、左側の列のセルに翻訳したい塊ごとに英文が入っている事が前提
        """

        # parent_path = kwargs['Fpath']
        # fname = kwargs['Fname'] + ".docx"
        # fpath = os.path.join(parent_path, fname)
        # docxfile = Document(fpath)
        # tables = self.docxfile.tables[0]
        column_eg = self.tables.columns[0]
        column_jp = self.tables.columns[1]                    
        sentences_eg = []
        #sentences_jp = []
        
        # print("translate start at row = ", self.insert_pt)
        for i in range(self.insert_pt, len(column_eg.cells)): # 未翻訳部分の英語から翻訳する
            sentences_eg.append(column_eg.cells[i].text)
            #sentences_jp.append(column_jp.cells[i])
        no_null_sentence = [text for text in sentences_eg if text != ""] # 空文字を消去する
        print(f"Total sensentces in this articl = {len(no_null_sentence)}")

        return no_null_sentence
            
    
    def _keyword_replace(self, sentence, type_of_dict: str="jp_to_jp"):
        
        """self.replace_dict_jp_to_jpで指定された単語変換を行う。2重変換が起きないようにマッチしたら直接変換するのでなく、
        辞書内の順番を示すインデックス値に1度変換してからstr.format()メソッドを使用して一斉に変換する。"""
        
        index_list = []
        replace_dict = self.replace_dict_jp_to_jp if type_of_dict == "jp_to_jp" else self.replace_dict_eg_to_jp

        "{ ←波カッコが文章中に含まれると最後の一括変換時にご認識されるので置換しておく"
        "具体的には、for文の1行面で、置換したいキーワードをそれが出てきた順番を波カッコで囲う形にしているのでそれに反応する。"
        sentence = sentence.replace("{", "$")
        sentence = sentence.replace("}", "!")

        for i, item in enumerate(replace_dict.items()):
            sentence = sentence.replace(item[0], '{'+str(i)+'}')
            index_list.append(item[1])
        sentence = sentence.format(*index_list)
        sentence = sentence.replace("$", "{")
        sentence = sentence.replace("!", "}")

        return sentence

    def _translation_deepl(self, sourse_texts=[]):

        """
        抜き出した英文をdeeplにコピペして翻訳する関数。引数は翻訳する文章が要素のリスト

        """

        # browser = webdriver.Chrome(ChromeDriverManager().install())

        # 2023/08/01現在、ChromeDriverManager()の引数にバージョンを指定しないとエラーが出るので、
        # 以下のようにして最新バージョンを取得して指定するように変更
        # 参考サイト(https://qiita.com/hs2023/questions/ffab105c5692692624ab)
        res = requests.get('https://chromedriver.storage.googleapis.com/LATEST_RELEASE')
        browser = webdriver.Chrome(ChromeDriverManager(res.text).install())
        url = 'https://www.deepl.com/ja/translator'
        browser.get(url)
        # deepleにアクセスするまでしばらく待つ
        time.sleep(1)
        # wordの文章をparagraph単位で翻訳していく
        # 先頭が数字の場合にセクションのタイトルと判定している。
        translated_texts = []
        for i in range(len(sourse_texts)):         
            sourse_text = sourse_texts[i]

            # セクション区切りをver4では廃止。単純にすべての英文が翻訳されるように変更(2022/07/01)
    #         is_numeric = re.match("\d", sourse_text)
    #         # else句以外はタイトルまたは数式判定のもので一致した場合はcontinueで次のループ要素へジャンプ
    #         if type(is_numeric) is re.Match:
    #             print("i={} continue\n: {}\n".format(i, sourse_text))
    #             translated_texts.append("")
    # #             print("continue\n")
    #             continue
    #         elif len(sourse_text) < 10: # 数式の場合は"Eq.1"など短い文章で示している想定なのでそれらをスキップ
    #             print("i={} continue\n: {}\n".format(i, sourse_text))
    #             translated_texts.append("")
    # #             print("continue\n")
    #             continue
    #         else:
            print("i={}\n: {}\n".format(i, sourse_text))

            stextarea = browser.find_element_by_css_selector(
                '.lmt__textarea.lmt__source_textarea.lmt__textarea_base_style')
            ttextarea = browser.find_element_by_css_selector(
                '.lmt__textarea.lmt__target_textarea.lmt__textarea_base_style')
            
            lock.acquire()
            # stextarea.send_keys(sourse_text) #この方法だと絵文字（数式を送る事ができない）
            # javascriptを仕込む方法に変更 2022/4/22 https://tech.bita.jp/article/19
            INPUT_EMOJI = """
            arguments[0].value += arguments[1];
            arguments[0].dispatchEvent(new Event('change'));
            """
            browser.execute_script(INPUT_EMOJI, stextarea, sourse_text)
            lock.release()
            #time.sleep(2)

            translated_text = ''

            lock.acquire()
            # 完全に翻訳されるまで繰り返す。判定は翻訳後文章の文字数が0かどうか
            # 文字数が0の場合は繰り返される。
            sec = self.wating_sec
            while not translated_text:
                # print("keep now....")
                time.sleep(sec)
                translated_text = ttextarea.get_property('value')
                print("translated_text: \n{}\n len: {}\n ".format(translated_text, len(translated_text))) 
                can_translated = re.search("\[\.\.\.\]", translated_text)
                if type(can_translated) is re.Match:
                    print("\ntry again\n")
                    translated_text = ''
                    sec += 2
                elif len(translated_text) == 0: # 翻訳されていない場合（[...]すら表示されていないパターン）
                    print("\ntry again\n")
                    translated_text = ''
                    sec += 2
                else:
                    if self.replace_flag_jp_to_jp:
                        translated_texts.append(self._keyword_replace(translated_text))
                    else:
                        translated_texts.append(translated_text)
            # ここに翻訳結果が出力されたら成功（途中経過をjupyter notebookで確認する場合はコメントイン）
            #print(translated_text)    
            lock.release()

            # sourse_textareaをクリアする
            time.sleep(2)
            stextarea.send_keys(Keys.CONTROL, "a")
            stextarea.send_keys(Keys.BACKSPACE)
                

        # print("finish translated")

        return translated_texts

    def _split_sentences(self, sentences=[]):

        """
        並列処理の前処理として翻訳する英文のリストを並列処理に使用する論理プロセッサの数で分割する関数
        """

        n = math.ceil(len(sentences)/self.max_worker)
        split_sentences = [sentences[i:i+n] for i in range(0, len(sentences), n)]

        return split_sentences

    def run_multiThread(self):

        """
        _translation_deepl関数を、引数辞書で指定した論理プロセッサ数で処理する。
        """

        # 翻訳前に専門用語（英語）を日本語に変換する（元の英文は変更しないように深いコピーをする）
        if self.replace_flag_eg_to_jp == True:
            sentences = deepcopy(self.eg_sentence)
            sentences = [self._keyword_replace(sentence, type_of_dict="eg_to_jp") for sentence in sentences]
        else:
            sentences = self.eg_sentence

        split_list = self._split_sentences(sentences=sentences)
        print("len(split_list): {}".format(len(split_list)))
        for i in range(len(split_list)):
            texts = split_list[i]
            print("i={}\nsplit_list: {}".format(i, texts))

        with ThreadPoolExecutor(max_workers=self.max_worker) as executor:

            result = executor.map(self._translation_deepl, split_list)

        # translated_texts = []
        results = [sentences for sentences in result]
        for i in range(len(results)):
            texts = results[i]
    #         print("len(translated_texts{}): {}\n".format(i, len(texts)))
            for j in range(len(texts)):
    #             print("i={} j={}\n{}".format(i, j, texts[j]))
                self.translated_sentence.append(texts[j])

        # return translated_texts


    def saveFile(self):

        # parent_path = kwargs['Fpath']
        # fname = kwargs['Fname'] + ".docx"
        # fpath = os.path.join(parent_path, fname)
        # docxfile = Document(fpath)
        # tables = docxfile.tables[0]
        o_cells = self.tables.column_cells(0) # original_text: english
        t_cells = self.tables.column_cells(1) # translated_text: japanese
    
        for i in range(self.insert_pt, len(self.eg_sentence)+self.insert_pt): # 未翻訳セルに書き込む
            # print(f"insert row = {i}")
            o_cells[i].text = self.eg_sentence[i-self.insert_pt]
            t_cells[i].text = self.translated_sentence[i-self.insert_pt]
            t_runs = t_cells[i].paragraphs[0].runs
            # 翻訳後の文章のフォントを変更
            for run in t_runs:
                run.font.name = self.japaneseFont
                run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
            
            # Ver4ではセクション区切りを廃止したため以下はコメントアウト（2022/07/01)
            # セクションを示す文字列なら文字のスタイルをボルドーにする。
            # 例： 1.Introduction -> 文字スタイルをボルドーに変換。
            # セクションかどうかの判定は先頭の文字が数字かどうかで判定している。
            # is_numeric = re.match("\d", self.eg_sentence[i-self.insert_pt])
            # if type(is_numeric) is re.Match:
            #     o_runs = o_cells[i].paragraphs[0].runs
            #     for run in o_runs:
            #         if run.text:
            #             run.bold = True


        self.docxfile.save(self.fpath)

    
if __name__ == "__main__":
    
    # 翻訳後に専門用語を正しく置換するための辞書オブジェクト
    word_dict_jp_to_jp = {
                "ニックスゾーン": "ニックゾーン",
                "ナイフゾーン": "ニックゾーン",
                "ニッケルゾーン": "ニックゾーン",
                "ニックスポイント": "ニックポイント",
                "キックポイント": "ニックポイント",
                "ナイフポイント": "ニックポイント",
                "ノッチポイント": "ニックポイント",
                "ノックポイント": "ニックポイント",
                "ニッケル点": "ニックポイント",
                "ニッケルポイント": "ニックポイント",
                "コツポイント": "ニックポイント",
                "刻み目": "ニックポイント",
                "浸食": "侵食",
                "切り込み": "侵食",
                "切開": "侵食",
                "切り下げ": "侵食",
                # "切り上げ": "侵食",
                "侵食効率": "侵食速度", 
                "侵食率": "侵食速度",
                "隆起率": "隆起速度",
                "侵食性": "侵食係数",
                "基盤レベル":"ベースレベル",
                "基底レベル":"ベ-スレベル",
                "基盤面": "ベースレベル",
                "ランドスケープ": "地形",
                "景観": "地形",
                "風景": "地形",
                "岩石の隆起": "隆起",
                "岩盤の隆起": "隆起",
                "捕獲": "河川争奪",
                "ストリームキャプチャー": "河川争奪",
                "海賊行為": "河川争奪",
                "集落": "流域",
                "集水域": "流域",
                "排水域": "流域",
                "排水面積": "流域面積",
                "排水路面積": "流域面積",
                "寄与面積": "流域面積",
                "救済": "レリーフ",
                "浮き上がり": "レリーフ",
                "盆地": "流域",
                "ディバイド" : "分水界",
                "分水嶺" : "分水界",
                "排水溝": "分水界",
                "流力モデル":"ストリームパワーモデル",
                "流線力法則":"ストリームパワーモデル",
                "流線力則":"ストリームパワーモデル",
                "流線力侵食モデル":"ストリームパワーモデル",
                "河道パワーインシジョンモデル":"ストリームパワーモデル",
                "います。" : "いる。",
                "排水路ネットワーク": "流路網",
                "排水ネットワーク": "流路網",
                "過渡状態" : "非平衡状態",
                "過渡的な" : "非平衡な",
                "幹線水路" : "本流",
                "幹線流" : "本流",
                "幹線河道" : "本流",
                "河道の険しさ": "急峻度指標",
                "河道の急勾配": "急峻度指標",
                "急勾配指標" : "急峻度指標",
                "水路の急勾配" : "急峻度指標",
                "前線":"前地",
                "オージュ":"造山帯",
                "編組河川":"網状河川",
                "塩水噴出粒":"掃流砂",
                "塩害":"サルテーション",
                "塩分":"サルテーション",
                "塩水化":"サルテーション化",
                "堤防":"側壁",
                "点棒":"ポントバー",
                "ベンド":"湾曲部",
                "正投影度":"蛇行度",
                "洞道率":"蛇行度",
                "洞道力":"蛇行度",
                "洞道度":"蛇行度",
                "洞窟度":"蛇行度",
                "正弦度":"蛇行度",
                "メアンダー":"蛇行",
                "洞窟":"蛇行",
                "現地観測":"フィールド調査",
                "現地調査":"フィールド調査",
                "水路図":"ハイドログラフ",
                "水力半径":"径深",
                "接水面積":"潤辺",
                "湿潤周囲":"潤辺",
                }
    
    word_dict_eg_to_jp = {
                "knickpoints" : "ニックポイント",
                "knickpoint" : "ニックポイント",
                "base level" : "ベースレベル",
                "relief" : "レリーフ",
                "sinuosity":"蛇行度",
                "knickzone" : "ニックゾーン",
                "abrasion" : "アブレージョン",
                "plucking" : "プラッキング",
                "reach scale": "リーチスケール", 
                "reach-scale": "リーチスケール", 
                "stream power": "ストリームパワー",
                "discharge" : "流量",
                "saltation" : "サルテーション",
                }
    
    kwargs = {"Fpath": r"F:\研究関連\文献リスト\Numerical Model", # 要約対象のwordファイルが格納されているディレクトリ
              "Fname": r"Coulthard et al 2006.docx", # wordファイル名（拡張子込みで指定する）
              "max_worker": 5, # 同時実行スレッド数。もし１つのスレッドで例外が発生した場合、デッドロックになってしまいエラーが伝搬してこないので、デバックする際は1スレッドに変更する
              "replace_dict_jp_to_jp" : word_dict_jp_to_jp,# 翻訳後に専門用語を正しく置換するための辞書(上記の辞書オブジェクト)
              "replace_dict_eg_to_jp" : word_dict_eg_to_jp, # 翻訳前に専門用語を日本語に変換してからDeeplに渡すことを想定
              "japaneseFont" : "游明朝 (本文のフォント - 日本語)", # 翻訳された日本語のフォントを指定（詳しいフォント名はwordを参照）
             } 
    
    translator = auto_translator(wating_sec=2, **kwargs)
    translator.run_multiThread()
    translator.saveFile()
    
